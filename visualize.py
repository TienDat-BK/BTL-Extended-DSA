import gradio as gr

import os
import pandas as pd
import docx

import re
from ftfy import fix_text

from HSmodule import *
from source.FaissSearch import *

from source.BloomDetection import *
from source.Preprocessor import *
from source.SimHashDetection import *
from source.minHashDetection import *
from sentence_transformers import SentenceTransformer


def read_file(filepath):
    #tach file extension
    ext = os.path.splitext(filepath)[1].lower()
    # tach du lieu tu file
    if ext == ".docx":
      doc = docx.Document(filepath)
      paragraphs = [(p.text.strip()) for p in doc.paragraphs if p.text.strip()]
    elif ext == ".txt":
      with open(filepath, "r", encoding="utf-8") as f:
        content = f.read()
      paragraphs = [p.strip() for p in content.split("\n\n") if p.strip()]
    elif ext == ".csv":
      table = pd.read_csv(filepath)
      data_cols = ["content", "text", "paragraph"]
      for col in data_cols:
        if col in table.columns:
           paragraphs = table[col].dropna().astype(str).tolist()
           break
    else:
        raise ValueError("Không tìm thấy cột 'text' / 'content' / 'paragraph' trong file CSV hoặc file không hợp lệ.")
    return paragraphs

modelSimHash = SimHashDetection()
modelMinHash = MinHashDetection()
modelBloomFaiss = BloomDetection()

def run_SimHash(paragraphs : list[str]) -> list[list[VectorRecord]]:
    doing = modelSimHash
    return doing.detect(paragraphs)


def run_Bloom_Sim_Faiss(paragraphs : list[str]) -> list[list[VectorRecord]]:
    doing = modelBloomFaiss
    return doing.detect(paragraphs)


def run_Min(paragraphs : list[str]) -> list[list[VectorRecord]]:
    doing = modelMinHash
    return doing.detect(paragraphs)


def representative_texts(l : list[list[VectorRecord]], p : list[str]) -> dict():
    text = {min(x.id for x in group) : max([p[x.id] for x in group], key=len) for group in l}
    return dict(sorted(text.items()))


def duplication_text(filepath : str, method : str):

    # tach du lieu
    paragraphs = read_file(filepath)
    if method == "SimHash (Semantic)":
      ans = run_SimHash(paragraphs)
    elif method == "Bloom + Faiss (Semantic)":
      ans = run_Bloom_Sim_Faiss(paragraphs)
    elif method == "MinHash (Syntax)":
      ans = run_Min(paragraphs)

    # loc van ban dai dien
    filted_text = representative_texts(ans, paragraphs)




    #hightlight
    # màu để highlight - 17 màu
    from docx.enum.text import WD_COLOR_INDEX
    colors = [
        None,  # 0 → không highlight
        WD_COLOR_INDEX.BLACK,
        WD_COLOR_INDEX.BLUE,
        WD_COLOR_INDEX.BRIGHT_GREEN,
        WD_COLOR_INDEX.DARK_BLUE,
        WD_COLOR_INDEX.DARK_RED,
        WD_COLOR_INDEX.DARK_YELLOW,
        WD_COLOR_INDEX.GRAY_25,
        WD_COLOR_INDEX.GRAY_50,
        WD_COLOR_INDEX.GREEN,
        WD_COLOR_INDEX.PINK,
        WD_COLOR_INDEX.RED,
        WD_COLOR_INDEX.TEAL,
        WD_COLOR_INDEX.TURQUOISE,
        WD_COLOR_INDEX.VIOLET,
        WD_COLOR_INDEX.WHITE,
        WD_COLOR_INDEX.YELLOW,
    ]

    # tạo dictionraty để truy xuất xử lý dữ liệu
    group_id = {x.id: (0 if len(group) == 1 else group_index) for group_index, group in enumerate(ans) for x in group}

    doc = docx.Document()
    doc.add_heading("Duplicate Text Highlighting Result")
    # Thêm từng paragraph từ filted_text
    for para in filted_text.values():
        doc.add_paragraph(para)

    # Lưu file
    result = "result.docx"
    doc.save(result)



    # in ket qua
    # thanh cuon
    css = """
    <style>
    .simple-scroll-content {
        max-height: 500px;
        overflow-y: auto;
        padding: 20px;
        border: 1px solid #444;
        border-radius: 10px;
        background: #000;
        color: white;
        margin-top: 10px;
    }
    </style>
    """

    html1 = css + """
    <div class="simple-scroll-content">
    """

    # chỉ hiển thị nhóm có 2 item trở lên
    cnt = 0
    for id, group in enumerate(ans):
        if len(group) < 2:
            continue
        cnt += 1
        html1 += f"<h3 style='color:#0af'>Group {cnt} ({len(group)} items)</h3><ul>"
        for para in group:
            html1 += f"<li><b>Paragraph {para.id}</b>: {paragraphs[para.id][:300]}...</li>"
        html1 += "</ul>"

    html1 += "</div>"


    # van ban sau khi loc
    html2 = css + """
    <div class="simple-scroll-content">
    """

    for para in filted_text.values():
      html2 += f"<h3 style='color:#0af'>{para}</h3>"

    html2 += "</div>"

    return html1, html2, result

# dùng gradio để tạo giao diện demo
with gr.Blocks(title="Duplicate Text Detector") as demo:
    gr.Markdown("## 🧩 Duplicate Text Detector")

    with gr.Row():
        with gr.Column():
            file_input = gr.File(label="Upload document file")
            method_choice = gr.Radio(
                ["SimHash (Semantic)", "Bloom + Faiss (Semantic)", "MinHash (Syntax)"],
                label="Choose method :",
                value="SimHash (Semantic)"
            )
            submit_btn = gr.Button("Submit", variant="primary")

    with gr.Row():
        with gr.Column():
            with gr.Accordion("📊 Duplication Text Group Result", open=False) as accordion_1:
              html_output_1 = gr.HTML()

        with gr.Column():
            with gr.Accordion("📝 Text after removing duplication", open=False) as accordion_2:
              html_output_2 = gr.HTML()

    with gr.Row():
        file_output = gr.File(label="Download result")

    submit_btn.click(
        fn=duplication_text,
        inputs=[file_input, method_choice],
        outputs=[html_output_1, html_output_2, file_output]
    )
    
port = int(os.environ.get("PORT", 7860))  # Render cung cấp PORT, fallback 7860 khi chạy local
demo.launch(server_name="0.0.0.0", server_port=port)
